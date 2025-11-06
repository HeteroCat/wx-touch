import docx

def extract_tables_from_doc():
    doc_path = r'C:\Users\h1930\Desktop\个人\AI应用开发\wx-touch\微信公众号爬虫服务 API 文档.docx'
    doc = docx.Document(doc_path)

    print("=== 提取表格内容 ===")
    for table_idx, table in enumerate(doc.tables):
        print(f"\n表格 {table_idx + 1}:")
        for row_idx, row in enumerate(table.rows):
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                print(f"  行 {row_idx + 1}: {' | '.join(row_text)}")

    # 也提取所有段落，寻找可能的代码
    print("\n=== 查找包含代码的段落 ===")
    code_keywords = ['import', 'requests', 'hashlib', 'md5', 'time', 'def', 'params']
    for para in doc.paragraphs:
        text = para.text.strip()
        if any(keyword in text for keyword in code_keywords):
            print(f"找到代码段落: {text}")

if __name__ == "__main__":
    extract_tables_from_doc()